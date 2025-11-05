import os
import shutil
import stat
import tempfile
import time
import json
import hashlib
import sqlite3

import pandas as pd

from typing import List, Dict, Optional, Callable, Any  # + Callable
from langchain_classic.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
# from langchain.vectorstores import Chroma
from langchain_classic.schema import Document
from langchain_community.document_loaders import TextLoader, DirectoryLoader
#from langchain_community.vectorstores import Chroma
from langchain_chroma import Chroma
from filelock import FileLock
from chromadb.config import Settings  # added
#from settings import settings  # добавлено

class VectorStore:
    def __init__(
        self,
        docs_path: str = "./data/docs/",
        vector_store_path: str = "./data/vector_store",
        openai_api_key: str =  os.getenv("OPENAI_API_KEY"),
        chunk_size: int = 5000,
        chunk_overlap: int = 500
    ):
        """Инициализация векторного хранилища
        
        Args:
            docs_path: путь к документам
            vector_store_path: путь для сохранения векторной БД
            openai_api_key: API ключ OpenAI
            chunk_size: размер чанка для разбиения текста
            chunk_overlap: перекрытие между чанками
        """
        # Используем переменные окружения или значения по умолчанию (централизованный fallback)
        self.docs_path = docs_path or os.getenv("DOCS_PATH")# or settings.DOCS_PATH
        self.vector_store_path = vector_store_path or os.getenv("VECTOR_STORE_PATH", "/app/data/vector_store")
        
        # Путь для файла блокировки
        self.lock_file_path = f"{self.vector_store_path}.lock"
        
        print(f"Текущая рабочая директория: {os.getcwd()}")

        # Преобразуем в абсолютные пути
        self.docs_path = os.path.abspath(self.docs_path)
        self.vector_store_path = os.path.abspath(self.vector_store_path)
        self.openai_api_key = openai_api_key or os.getenv("OPENAI_API_KEY")
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # Прогресс-колбек (опционально)
        self.progress_callback: Optional[Callable[[str], None]] = None

        # Параметры обработки Excel
        self.excel_max_rows = int(os.getenv("EXCEL_MAX_ROWS", "2000"))
        self.excel_output_format = (
            os.getenv("EXCEL_NORMALIZED_FORMAT", "markdown") or "markdown"
        ).strip().lower()

        # Сохранение копий исходных документов
        self.save_normalized_docs = str(os.getenv("SAVE_NORMALIZED_DOCS", "false")).lower() in ("1", "true", "yes")
        self.save_only_changed = str(os.getenv("SAVE_ONLY_CHANGED", "true")).lower() in ("1", "true", "yes")
        normalized_root = os.getenv("NORMALIZED_OUTPUT_DIR", "./data/normalized")
        self.normalized_base_dir = os.path.abspath(normalized_root)
        self.normalized_output_dir = self._resolve_normalized_output_dir()
        self.include_normalized_docs = str(
            os.getenv("INCLUDE_NORMALIZED_DOCS", "true")
        ).lower() in ("1", "true", "yes")
        if self.save_normalized_docs or self.include_normalized_docs:
            try:
                os.makedirs(self.normalized_output_dir, exist_ok=True)
            except Exception as e:
                msg = (
                    "Предупреждение: не удалось создать директорию для нормализованных файлов: "
                    f"{e}"
                )
                if self.save_normalized_docs:
                    raise PermissionError(msg) from e
                print(msg)

        # Явно отключаем анонимную телеметрию Chroma (во избежание сбоев в CI)
        os.environ.setdefault("ANONYMIZED_TELEMETRY", "False")
        self.client_settings = Settings(
            anonymized_telemetry=False,
            allow_reset=True,
            is_persistent=True,
        )

        if not self.openai_api_key:
            raise ValueError("OPENAI_API_KEY не установлен")

        # Инициализация эмбеддингов
        self.embeddings = OpenAIEmbeddings(
            model="text-embedding-3-large",
            openai_api_key=self.openai_api_key
        )

        # Инициализация сплиттера
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            is_separator_regex=False
        )

        self.vectorstore = None

        # Создаем директории с правильными правами
        self._ensure_directories()

    def set_progress_callback(self, cb: Optional[Callable[[str], None]]):
        """Устанавливает колбек прогресса (строка → уведомление)."""
        self.progress_callback = cb

    def _notify(self, msg: str):
        """Локальная отправка уведомления: печать + внешний колбек (если задан)."""
        try:
            print(msg)
        finally:
            if self.progress_callback:
                try:
                    self.progress_callback(msg)
                except Exception:
                    # Никогда не валим процесс из-за колбека прогресса
                    pass

    def _resolve_normalized_output_dir(self) -> str:
        """Определяет каталог для нормализованных файлов текущего продукта."""
        docs_root_env = os.getenv("DOCS_PATH")# or settings.DOCS_PATH
        docs_abs = self.docs_path
        normalized_rel: Optional[str] = None

        if docs_root_env:
            docs_root_abs = os.path.abspath(docs_root_env)
            try:
                rel_candidate = os.path.relpath(docs_abs, start=docs_root_abs)
                if not rel_candidate.startswith(".."):
                    normalized_rel = rel_candidate if rel_candidate != "." else ""
            except Exception:
                normalized_rel = None

        if not normalized_rel:
            basename = os.path.basename(docs_abs.rstrip(os.sep))
            normalized_rel = basename or "product"

        return os.path.join(self.normalized_base_dir, normalized_rel)

    def _rel_path_from_docs(self, src: str) -> str:
        """Возвращает относительный путь от каталога docs до файла src."""
        try:
            return os.path.relpath(src, start=self.docs_path)
        except Exception:
            return os.path.basename(src) if src else "unknown"

    def _build_norm_path(self, doc: Document, changed: bool, idx: int, force: bool = False) -> Optional[str]:
        """Строит путь для сохранения нормализованного текста, учитывая Excel-листы и зеркальную структуру.
        Возвращает None, если сохранять не нужно (save_only_changed и нет изменений),
        за исключением случаев, когда требуется принудительное сохранение (force=True).
        """
        if self.save_only_changed and not changed and not force:
            return None
        src = doc.metadata.get("source")
        rel = self._rel_path_from_docs(src) if src else f"unknown_{idx}.txt"
        rel_dir = os.path.dirname(rel)
        base_name = os.path.basename(rel)
        # Для Excel добавляем суффикс с именем листа
        if doc.metadata.get("excel"):
            sheet = str(doc.metadata.get("sheet_name", "Sheet")).replace("/", "_").replace("\\", "_")
            out_name = f"{base_name}__{sheet}.norm.txt"
        else:
            out_name = f"{base_name}.norm.txt"
        out_dir = os.path.join(self.normalized_output_dir, rel_dir)
        try:
            os.makedirs(out_dir, exist_ok=True)
        except Exception:
            pass
        return os.path.join(out_dir, out_name)

    def _save_normalized_copy(
        self,
        doc: Document,
        new_text: str,
        changed: bool,
        idx: int,
        force: bool = False,
    ) -> Optional[str]:
        """Сохраняет текст на диск (нормализованный или исходный), возвращает путь или None."""
        if not self.save_normalized_docs:
            return None
        path = self._build_norm_path(doc, changed, idx, force=force)
        if not path:
            return None
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(new_text)
            try:
                rel = os.path.relpath(path, start=self.normalized_output_dir)
            except Exception:
                rel = path
            self._notify(f"[norm] 💾 Сохранено: {rel}")
            return path
        except Exception as e:
            print(f"Предупреждение: не удалось сохранить нормализованный файл '{path}': {e}")
            return None

    def _save_raw_documents(self, documents: List[Document]) -> None:
        """Опционально сохраняет неизмённые копии документов на диск."""
        if not self.save_normalized_docs:
            return

        self._notify("💾 Сохранение копий исходных документов")
        processed = 0
        saved_files = 0
        start_all = time.time()
        for doc in documents:
            try:
                content = doc.page_content if doc.page_content is not None else ""
                if content == "":
                    continue
                processed += 1
                src = doc.metadata.get("source", "unknown")
                fname = os.path.basename(src) if src and src != "unknown" else f"doc_{processed}"
                self._notify(f"[norm] ▶️ {fname}: сохраняем оригинальный текст")

                saved_path = self._save_normalized_copy(
                    doc,
                    new_text=content,
                    changed=False,
                    idx=processed,
                    force=True,
                )
                if saved_path:
                    saved_files += 1
            except Exception as e:
                print(
                    f"Предупреждение: пропущен документ при сохранении без изменений"
                    f" ({doc.metadata.get('source', 'unknown')}): {e}"
                )
        total_time = time.time() - start_all
        self._notify(
            f"✅ Обработано документов: {processed} | Сохранено файлов: {saved_files}"
            f" | время всего={total_time:.2f}s"
        )

    def _ensure_directories(self):
        """Создание необходимых директорий"""
        try:
            os.makedirs(self.docs_path, exist_ok=True)
        except OSError as exc:
            raise PermissionError(
                f"Не удалось создать директорию документов {self.docs_path}: {exc}"
            ) from exc

        try:
            os.makedirs(self.vector_store_path, exist_ok=True)
        except OSError as exc:
            raise PermissionError(
                f"Не удалось создать директорию векторного хранилища {self.vector_store_path}: {exc}"
            ) from exc

        # Проверяем, что хранилище и его родитель доступны для записи (обнаруживаем read-only маунты заранее)
        self._assert_writable_dir(
            self.vector_store_path,
            reason="запись в директорию векторного хранилища",
        )
        parent_dir = os.path.dirname(self.vector_store_path) or self.vector_store_path
        self._assert_writable_dir(
            parent_dir,
            reason="создание временных директорий для пересборки",
        )

        if self.save_normalized_docs:
            self._assert_writable_dir(
                self.normalized_output_dir,
                reason="сохранение нормализованных файлов",
            )

    def _is_learning_source(self, source: Optional[str]) -> bool:
        """Определяет, относится ли путь к файлу режима обучения (начинается с learn)."""
        if not source:
            return False
        base_name = os.path.basename(source).lower()
        if base_name.startswith("learn"):
            return True
        parent = os.path.basename(os.path.dirname(source)).lower()
        return parent.startswith("learn")

    def _mark_learning_documents(self, documents: List[Document]) -> None:
        """Помечает документы режима обучения в метаданных."""
        for doc in documents:
            source = doc.metadata.get("source")
            if self._is_learning_source(source):
                doc.metadata["learning_doc"] = True
                doc.metadata.setdefault("priority", "learn")

    def _doc_priority_rank(self, doc: Document) -> int:
        """Возвращает числовой приоритет документа (меньше = важнее)."""
        if doc.metadata.get("learning_doc"):
            return 0
        if doc.metadata.get("normalized"):
            return 1
        return 2

    def _sort_documents_by_priority(self, documents: List[Document]) -> List[Document]:
        """Сортирует документы с учётом приоритетов (режим обучения выше всего)."""
        return sorted(
            documents,
            key=lambda doc: (self._doc_priority_rank(doc), doc.metadata.get("source", ""))
        )

    def _assert_writable_dir(self, directory: str, reason: str = "") -> None:
        """Проверяет возможность записи в директорию, создавая и удаляя временный файл."""
        try:
            os.makedirs(directory, exist_ok=True)
        except OSError as exc:
            raise PermissionError(
                f"Не удалось создать директорию {directory}: {exc}"
            ) from exc

        hint = f" ({reason})" if reason else ""
        try:
            fd, tmp_path = tempfile.mkstemp(prefix=".writetest_", dir=directory)
            os.close(fd)
            os.unlink(tmp_path)
        except OSError as exc:
            raise PermissionError(
                f"Нет прав на запись в директорию {directory}{hint}: {exc}"
            ) from exc

    # === Excel helpers ===
    def _sanitize_markdown_cell(self, value: Optional[str]) -> str:
        """Подготавливает значение ячейки к безопасному отображению в Markdown-таблице."""
        if value is None:
            return ""
        sanitized = str(value)
        sanitized = sanitized.replace("\r\n", "\n").replace("\r", "\n").strip()
        sanitized = sanitized.replace("|", "\\|")
        return sanitized.replace("\n", "<br>")

    def _df_to_markdown(self, df: "pd.DataFrame") -> str:
        """Формирует Markdown-таблицу из DataFrame без внешних зависимостей."""
        safe_df = df.fillna("")
        headers = [
            self._sanitize_markdown_cell(str(col) if str(col) else f"Column {idx + 1}")
            for idx, col in enumerate(safe_df.columns)
        ]
        lines = [
            "| " + " | ".join(headers) + " |",
            "| " + " | ".join("---" for _ in headers) + " |",
        ]
        for _, row in safe_df.iterrows():
            cells = [self._sanitize_markdown_cell(row[col]) for col in safe_df.columns]
            lines.append("| " + " | ".join(cells) + " |")
        return "\n".join(lines)

    def _df_to_bullet(self, df: "pd.DataFrame") -> str:
        """Формирует компактное описание строк таблицы в виде маркированного списка."""
        safe_df = df.fillna("")
        records = safe_df.to_dict(orient="records")
        lines = []
        for idx, record in enumerate(records, start=1):
            lines.append(f"Строка {idx}:")
            for column in safe_df.columns:
                value = str(record.get(column, "")).strip()
                if not value:
                    continue
                value = value.replace("\r\n", " ").replace("\n", " ")
                lines.append(f"  - {column}: {value}")
            lines.append("")
        return "\n".join(lines).strip()

    def _df_to_tsv(self, df: "pd.DataFrame") -> str:
        """Возвращает TSV-представление таблицы (исторический формат по умолчанию)."""
        return df.to_csv(sep="\t", index=False)

    def _format_excel_table(self, df: "pd.DataFrame") -> (str, str):
        """Форматирует таблицу согласно настройкам и возвращает текст + применённый формат."""
        requested = (self.excel_output_format or "markdown").lower()
        aliases = {
            "md": "markdown",
            "table": "markdown",
            "list": "bullet",
        }
        fmt = aliases.get(requested, requested)

        try:
            if fmt == "markdown":
                return self._df_to_markdown(df), "markdown"
            if fmt == "bullet":
                return self._df_to_bullet(df), "bullet"
            if fmt == "csv":
                return df.to_csv(index=False), "csv"
            if fmt == "tsv":
                return self._df_to_tsv(df), "tsv"
            self._notify(f"[norm] ⚠️ Неизвестный формат Excel '{requested}', используем TSV")
            return self._df_to_tsv(df), "tsv"
        except Exception as e:
            self._notify(
                f"[norm] ⚠️ Ошибка форматирования Excel ({fmt}): {e}. Используем TSV"
            )
            return self._df_to_tsv(df), "tsv"

    def _excel_to_text(self, file_path: str) -> List[Document]:
        """Конвертирует .xlsx в список Document (по одному на лист).
        Текст формируется как простой CSV/TSV без индексов, ограничиваем кол-во строк.
        """
        docs: List[Document] = []
        try:
            import pandas as pd  # импорт внутри функции, чтобы не падать без pandas
        except Exception as e:
            print(f"Предупреждение: pandas не установлен, пропускаем Excel: {e}")
            return docs

        try:
            # Читаем все листы как строки (для устойчивости к типам)
            xls = pd.read_excel(file_path, sheet_name=None, dtype=str, engine="openpyxl")
            total_sheets = len(xls)
            self._notify(f"Загрузка Excel: {os.path.basename(file_path)} | листов: {total_sheets}")
            for sheet_name, df in xls.items():
                if df is None:
                    continue
                # Нормализуем NaN
                df = df.fillna("")
                # Ограничим строки, чтобы не раздувать документ
                max_rows = self.excel_max_rows
                rows_before = len(df)
                if rows_before > max_rows:
                    df = df.head(max_rows)
                table_text, used_format = self._format_excel_table(df)
                format_label = {
                    "markdown": "Markdown",
                    "bullet": "bullet list",
                    "csv": "CSV",
                    "tsv": "TSV",
                }.get(used_format, used_format)
                header = (
                    f"Источник: {os.path.basename(file_path)} | Лист: {sheet_name} | Строк: {len(df)}"
                    + (f" (обрезано из {rows_before})" if rows_before > len(df) else "")
                    + f" | Формат таблицы: {format_label}"
                )
                full_text = header + "\n\n" + table_text
                docs.append(
                    Document(
                        page_content=full_text,
                        metadata={
                            "source": file_path,
                            "sheet_name": str(sheet_name),
                            "rows": len(df),
                            "excel": True,
                        },
                    )
                )
            return docs
        except Exception as e:
            print(f"Предупреждение: ошибка чтения Excel '{file_path}': {e}")
            return docs

    def _load_excel_documents(self) -> List[Document]:
        """Поиск и загрузка .xlsx из self.docs_path рекурсивно."""
        excel_docs: List[Document] = []
        try:
            for root, _, files in os.walk(self.docs_path):
                for name in files:
                    # Поддерживаем только .xlsx (без .xls, чтобы не тянуть xlrd)
                    if name.lower().endswith(".xlsx"):
                        path = os.path.join(root, name)
                        docs = self._excel_to_text(path)
                        excel_docs.extend(docs)
        except Exception as e:
            print(f"Предупреждение: ошибка при поиске Excel: {e}")
        self._notify(f"Excel: сформировано {len(excel_docs)} документ(ов) из .xlsx файлов")
        return excel_docs

    def _clear_vector_store(self):
        """Очистка векторного хранилища"""
        if os.path.exists(self.vector_store_path):
            shutil.rmtree(self.vector_store_path)

    def _load_normalized_documents(self) -> List[Document]:
        """Загрузка нормализованных файлов для индексации."""
        if not self.include_normalized_docs:
            return []
        if not os.path.isdir(self.normalized_output_dir):
            return []

        loader = DirectoryLoader(
            self.normalized_output_dir,
            glob="**/*.norm.txt",
            loader_cls=TextLoader,
            loader_kwargs={'encoding': 'utf-8'},
        )

        try:
            docs = loader.load()
        except Exception as e:
            print(f"Предупреждение: ошибка при загрузке нормализованных файлов: {e}")
            return []

        for doc in docs:
            doc.metadata["normalized"] = True
        return docs

    def _shutdown_chroma_client(self, store: Any) -> None:
        """Attempt to stop embedded Chroma client to release file handles."""
        if store is None:
            return
        try:
            client = getattr(store, "_client", None)
            system = getattr(client, "_system", None)
            stop = getattr(system, "stop", None)
            if callable(stop):
                stop()
        except Exception as exc:
            print(f"[chroma] Failed to stop client cleanly: {exc}")

    def _force_close_connections(self):
        """Принудительное закрытие всех соединений ChromaDB"""
        try:
            import gc
            import threading
            import subprocess
            
            # Закрываем текущее соединение если есть (без внутренних reset/stop)
            if hasattr(self, 'vectorstore') and self.vectorstore is not None:
                try:
                    self._shutdown_chroma_client(self.vectorstore)
                except Exception as stop_error:
                    print(f"[chroma] Could not stop active store cleanly: {stop_error}")
                try:
                    # Удаляем ссылку на объект — этого достаточно для локального клиента
                    del self.vectorstore
                    self.vectorstore = None
                except Exception as e:
                    print(f"Предупреждение при закрытии vectorstore: {e}")
            
            # Принудительная сборка мусора
            gc.collect()
            
            # Принудительное освобождение файлов через lsof
            try:
                # Ищем процессы, использующие файлы в директории векторного хранилища
                result = subprocess.run(
                    ['lsof', '+D', self.vector_store_path],
                    capture_output=True, text=True, timeout=10
                )
                if result.returncode == 0 and result.stdout.strip():
                    print(f"Найдены процессы, использующие файлы: {result.stdout}")
                    # Просто даем больше времени на освобождение
                    time.sleep(3)
                else:
                    print("Файлы векторного хранилища не используются другими процессами")
            except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
                # lsof может быть недоступен или таймаут
                pass
            
            # Даем время системе освободить ресурсы
            time.sleep(3)
            
            print("Соединения ChromaDB закрыты")
            
        except Exception as e:
            print(f"Ошибка при принудительном закрытии соединений: {e}")

    def _safe_remove_directory(self, directory_path: str, max_attempts: int = 5):
        """Безопасное удаление директории с повторными попытками"""
        import subprocess
        
        for attempt in range(max_attempts):
            try:
                if os.path.exists(directory_path):
                    # Сначала пытаемся изменить права доступа для всех файлов
                    for root, dirs, files in os.walk(directory_path):
                        for file in files:
                            file_path = os.path.join(root, file)
                            try:
                                os.chmod(file_path, stat.S_IWRITE | stat.S_IREAD)
                            except:
                                pass
                    
                    # Удаляем директорию
                    shutil.rmtree(directory_path)
                    print(f"Директория {directory_path} успешно удалена")
                    return
            except OSError as e:
                if e.errno == 16 and attempt < max_attempts - 1:  # Device or resource busy
                    print(f"Директория занята, попытка {attempt + 1}/{max_attempts}. Ждем...")
                    
                    # Дополнительные попытки освобождения ресурсов
                    if attempt >= 2:  # После второй попытки становимся более агрессивными
                        try:
                            # Попробуем принудительно закрыть процессы, использующие файлы
                            result = subprocess.run(
                                ['lsof', '+D', directory_path],
                                capture_output=True, text=True, timeout=5
                            )
                            if result.returncode == 0 and result.stdout.strip():
                                print("Обнаружены процессы, использующие файлы векторного хранилища")
                        except:
                            pass
                    
                    # Увеличиваем время ожидания с каждой попыткой
                    wait_time = min(2 ** attempt, 10)
                    time.sleep(wait_time)
                    continue
                elif attempt == max_attempts - 1:
                    print(f"Не удалось удалить директорию после {max_attempts} попыток: {e}")
                    # В крайнем случае, переименуем директорию, чтобы не блокировать работу
                    try:
                        backup_path = f"{directory_path}_backup_{int(time.time())}"
                        os.rename(directory_path, backup_path)
                        print(f"Директория переименована в {backup_path} для последующего удаления")
                        return
                    except:
                        raise e
                else:
                    raise

    def _load_documents(self) -> List[Document]:
        print("\n===== ФАЙЛОВАЯ СИСТЕМА =====")
        print(f"Путь к документам: {self.docs_path}")

        # Проверка существования директории
        if not os.path.exists(self.docs_path):
            print(f"ОШИБКА: Директория {self.docs_path} не существует!")
            return []

        # Проверка содержимого директории
        files = os.listdir(self.docs_path)
        print(f"Файлы в директории: {files}")

        # Загрузка документов нескольких типов
        print("Загрузка документов...")
        documents: List[Document] = []

        # TXT
        txt_loader = DirectoryLoader(
            self.docs_path,
            glob="**/*.txt",
            loader_cls=TextLoader,
            loader_kwargs={'encoding': 'utf-8'}
        )
        txt_docs = []
        try:
            txt_docs = txt_loader.load()
            documents.extend(txt_docs)
        except Exception as e:
            print(f"Предупреждение: ошибка при загрузке .txt: {e}")

        # DOCX (через python-docx)
        docx_docs = []
        try:
            from docx import Document
            from langchain_classic.schema import Document as LangchainDocument
            
            for root, dirs, files in os.walk(self.docs_path):
                for file in files:
                    if file.endswith('.docx'):
                        file_path = os.path.join(root, file)
                        try:
                            doc = Document(file_path)
                            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                            if text.strip():
                                docx_docs.append(LangchainDocument(
                                    page_content=text,
                                    metadata={"source": file_path}
                                ))
                        except Exception as file_e:
                            print(f"Error loading file {file_path}: {file_e}")
            documents.extend(docx_docs)
        except Exception as e:
            print(f"Предупреждение: не удалось загрузить .docx (установите зависимости python-docx): {e}")

        # PDF (через unstructured)
        pdf_docs = []
        try:
            from langchain_community.document_loaders import UnstructuredPDFLoader
            pdf_loader = DirectoryLoader(
                self.docs_path,
                glob="**/*.pdf",
                loader_cls=UnstructuredPDFLoader,
            )
            pdf_docs = pdf_loader.load()
            documents.extend(pdf_docs)
        except Exception as e:
            print(f"Предупреждение: не удалось загрузить .pdf (установите зависимости pypdf/unstructured): {e}")

        # XLSX (через pandas → текст)
        excel_docs = []
        try:
            excel_docs = self._load_excel_documents()
            documents.extend(excel_docs)
        except Exception as e:
            print(f"Предупреждение: ошибка при обработке Excel: {e}")

        normalized_docs: List[Document] = []
        try:
            normalized_docs = self._load_normalized_documents()
            documents.extend(normalized_docs)
        except Exception as e:
            print(f"Предупреждение: ошибка при добавлении нормализованных документов: {e}")

        # Помечаем документы режима обучения до дедупликации
        self._mark_learning_documents(documents)

        unique_docs: List[Document] = []
        index_by_hash: Dict[str, int] = {}
        for doc in documents:
            content = doc.page_content or ""
            content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()
            if content_hash in index_by_hash:
                existing_idx = index_by_hash[content_hash]
                existing_doc = unique_docs[existing_idx]
                if self._doc_priority_rank(doc) < self._doc_priority_rank(existing_doc):
                    unique_docs[existing_idx] = doc
                continue
            index_by_hash[content_hash] = len(unique_docs)
            unique_docs.append(doc)

        deduplicated = len(documents) - len(unique_docs)
        if deduplicated:
            self._notify(f"Удалено дубликатов документов: {deduplicated}")

        documents = self._sort_documents_by_priority(unique_docs)

        self._notify(
            f"Загружено {len(documents)} документов (txt: {len(txt_docs)}, docx: {len(docx_docs)}, pdf: {len(pdf_docs)}, "
            f"excel: {len(excel_docs)}, normalized: {len(normalized_docs)})"
        )
        return documents

    def _split_documents(self, documents: List[Document]) -> List[Document]:
        """Разбиение документов на чанки"""
        chunks = self.text_splitter.split_documents(documents)
        print(f"Создано {len(chunks)} чанков")
        return chunks

    def create(self, force_recreate: bool = True) -> 'VectorStore':
        """Создание или пересборка векторного хранилища."""
        self._notify("Создание нового векторного хранилища...")

        # Обновляем директории и проверяем права
        self._ensure_directories()

        if not force_recreate:
            try:
                if not self.is_empty():
                    self._notify(
                        "Хранилище уже существует — выполняем загрузку (force_recreate=False)"
                    )
                    return self.load()
            except Exception:
                # Если не удалось загрузить, продолжаем с пересборкой
                pass

        # Safe rebuild сохраняет текущий индекс до успешной пересборки
        return self.safe_rebuild()

    def load(self) -> 'VectorStore':
        """Загрузка существующего векторного хранилища"""
        if not os.path.exists(self.vector_store_path):
            raise ValueError(f"Векторное хранилище не найдено в {self.vector_store_path}")

        if not os.access(self.vector_store_path, os.W_OK):
            raise PermissionError(f"Нет прав на запись в директорию {self.vector_store_path}")

        self.vectorstore = Chroma(
            persist_directory=self.vector_store_path,
            embedding_function=self.embeddings,
            client_settings=self.client_settings,
        )

        print(f"Загружено векторное хранилище")
        return self

    def search(self, query: str, n_results: int = 6, product: str = "default") -> List[Dict]:
        """Поиск похожих документов

        Args:
            query: поисковый запрос
            n_results: количество результатов

        Returns:
            List[Dict]: список найденных документов
        """
        print(self.vectorstore)
        
        # Проверяем, что векторное хранилище инициализировано
        if not self.vectorstore:
            print("Векторное хранилище не инициализировано, пытаемся загрузить...")
            try:
                self.load()
            except Exception as e:
                print(f"Не удалось загрузить векторное хранилище: {e}")
                return []

        try:
            search_kwargs = None
            if product != "default":
                search_kwargs={"product": {"$eq": product}}
            search_result = self.vectorstore.similarity_search_with_relevance_scores(
                query,
                k=n_results,
                filter=search_kwargs
            )
            docs = [Document(page_content=d.page_content, metadata={**(d.metadata or {}), "relevance_score": float(s)}) for d, s in search_result if (s or 0) > 0]
            if not docs:
                return []

            # На случай старых индексов помечаем learning документы по их источнику
            self._mark_learning_documents(docs)

            prioritized = self._sort_documents_by_priority(docs)
            return prioritized[:n_results]
        except Exception as e:
            print(f"Ошибка при поиске в векторном хранилище: {e}")
            return []

    def safe_rebuild(self) -> 'VectorStore':
        """Безопасная пересборка векторного хранилища через переключение директорий"""
        self._notify("Начинается безопасная пересборка векторного хранилища...")
        # Повторно удостоверяемся, что можем писать в директории (могли смениться права или маунт)
        self._ensure_directories()

        # Создаем блокировку для предотвращения одновременного доступа
        with FileLock(self.lock_file_path, timeout=30):
            try:
                # Создаем новую директорию для векторного хранилища
                new_store_path = f"{self.vector_store_path}_new_{int(time.time())}"
                self._notify(f"Создается новое хранилище в: {new_store_path}")
                
                # Загружаем и обрабатываем документы
                documents = self._load_documents()
                if not documents:
                    raise ValueError("Не найдено документов для обработки")

                # Сохраняем копии исходных документов при необходимости
                self._save_raw_documents(documents)

                chunks = self._split_documents(documents)
                if not chunks:
                    raise ValueError("Не удалось создать чанки из документов")

                # Создаем новое хранилище в новой директории
                self._notify("Создание нового векторного хранилища...")
                new_vectorstore = Chroma.from_documents(
                    documents=chunks,
                    embedding=self.embeddings,
                    persist_directory=new_store_path,
                    client_settings=self.client_settings,
                )
                
                # Сохраняем новое хранилище
                #new_vectorstore.persist()
                self._notify("Новое хранилище создано и сохранено успешно")
                
                # Закрываем новое временное соединение
                self._shutdown_chroma_client(new_vectorstore)
                del new_vectorstore
                
                # Закрываем старое хранилище
                if self.vectorstore:
                    self._notify("Закрываем старое хранилище...")
                    self._force_close_connections()
                
                # Создаем backup старого хранилища если оно существует
                backup_path = None
                if os.path.exists(self.vector_store_path):
                    backup_path = f"{self.vector_store_path}_backup_{int(time.time())}"
                    self._notify(f"Создаем backup старого хранилища: {backup_path}")
                    try:
                        shutil.move(self.vector_store_path, backup_path)
                    except Exception as e:
                        self._notify(f"Не удалось создать backup: {e}")
                        # Если не можем переместить, просто продолжаем
                        backup_path = None
                
                # Перемещаем новое хранилище на место старого
                self._notify(f"Активируем новое хранилище...")
                try:
                    shutil.move(new_store_path, self.vector_store_path)
                    self._notify("Новое хранилище успешно активировано")
                except Exception as e:
                    # Если не удалось переместить, восстанавливаем backup
                    self._notify(f"Ошибка при активации нового хранилища: {e}")
                    if backup_path and os.path.exists(backup_path):
                        self._notify("Восстанавливаем backup...")
                        shutil.move(backup_path, self.vector_store_path)
                        backup_path = None
                    raise
                
                # Загружаем новое хранилище
                self.load()
                self._notify("Пересборка векторного хранилища завершена успешно!")
                
                # Удаляем backup если все прошло успешно
                if backup_path and os.path.exists(backup_path):
                    try:
                        self._notify(f"Удаляем backup: {backup_path}")
                        shutil.rmtree(backup_path)
                    except Exception as e:
                        self._notify(f"Предупреждение: не удалось удалить backup {backup_path}: {e}")
                
                # Сохранение манифеста
                self._save_manifest(self._gather_source_files())
                
                return self
                
            except Exception as e:
                # Очищаем временную директорию в случае ошибки
                if 'new_store_path' in locals() and os.path.exists(new_store_path):
                    try:
                        shutil.rmtree(new_store_path)
                        print(f"Временная директория {new_store_path} очищена")
                    except Exception as cleanup_error:
                        print(f"Не удалось очистить временную директорию: {cleanup_error}")
                # Если ошибка связана с попыткой записи в read-only SQLite, преобразуем её в PermissionError
                if isinstance(e, sqlite3.OperationalError) and "readonly" in str(e).lower():
                    readonly_msg = (
                        "Хранилище недоступно для записи (read-only). "
                        "Проверьте права на каталог или настройки маунта volume."
                    )
                    self._notify(f"Ошибка при пересборке хранилища: {readonly_msg}")
                    raise PermissionError(readonly_msg) from e

                self._notify(f"Ошибка при пересборке хранилища: {e}")
                raise
                
    def is_empty(self) -> bool:
        """Проверка, пусто ли векторное хранилище (файловая проверка без запуска Chroma)"""
        if not os.path.exists(self.vector_store_path):
            return True
            
        try:
            files = os.listdir(self.vector_store_path)
            if len(files) == 0:
                return True

            # Минимальный набор файлов Chroma (sqlite-хранилище)
            sqlite_path = os.path.join(self.vector_store_path, 'chroma.sqlite3')
            has_sqlite = os.path.isfile(sqlite_path) and os.path.getsize(sqlite_path) > 0

            # Часто рядом с БД есть директория коллекции (UUID)
            has_any_dir = any(
                os.path.isdir(os.path.join(self.vector_store_path, f)) for f in files
            )

            if has_sqlite:
                print(f"Векторное хранилище содержит файлы: {files}")
                return False

            # Если нет sqlite, но есть директория — все равно считаем пустым/битым
            if not has_sqlite and has_any_dir:
                print("Векторное хранилище существует, но не содержит валидной БД chroma.sqlite3")
                return True

            return True
        except Exception as e:
            print(f"Ошибка при доступе к директории хранилища: {e}")
            return True

    # ================== МАНИФЕСТ / ПРОВЕРКА АКТУАЛЬНОСТИ ==================
    @property
    def _manifest_path(self) -> str:
        return os.path.join(self.vector_store_path, "_manifest.json")

    def _gather_source_files(self) -> List[Dict[str, str]]:
        """Сканирует self.docs_path и собирает список файлов-источников с mtime и размером.
        Поддерживаемые расширения: .txt, .docx, .xlsx, .pdf
        """
        exts = {".txt", ".docx", ".xlsx", ".pdf"}
        results: List[Dict[str, str]] = []
        if not os.path.isdir(self.docs_path):
            return results
        for root, _, files in os.walk(self.docs_path):
            for f in files:
                _, ext = os.path.splitext(f.lower())
                if ext in exts:
                    full = os.path.join(root, f)
                    try:
                        st = os.stat(full)
                        rel = os.path.relpath(full, self.docs_path)
                        results.append({
                            "rel": rel,
                            "mtime": str(int(st.st_mtime)),
                            "size": str(st.st_size),
                        })
                    except Exception:
                        pass
        # Сортируем для детерминированного хеша
        results.sort(key=lambda x: x["rel"])
        return results

    def _compute_signature(self, items: List[Dict[str, str]]) -> str:
        h = hashlib.sha256()
        for it in items:
            h.update(it["rel"].encode("utf-8"))
            h.update(b"|")
            h.update(it["mtime"].encode("utf-8"))
            h.update(b"|")
            h.update(it["size"].encode("utf-8"))
            h.update(b"\n")
        return h.hexdigest()

    def _load_manifest(self) -> Optional[Dict]:
        try:
            if os.path.isfile(self._manifest_path):
                with open(self._manifest_path, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception:
            return None
        return None

    def _save_manifest(self, items: List[Dict[str, str]]):
        data = {
            "files": items,
            "signature": self._compute_signature(items),
            "timestamp": int(time.time()),
        }
        try:
            with open(self._manifest_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Предупреждение: не удалось сохранить манифест: {e}")

    def needs_rebuild(self) -> bool:
        """Проверяет, требует ли индекс пересборки.
        Критерии:
        - Отсутствует директория хранилища или chroma.sqlite3
        - Отсутствует манифест
        - Подпись файлов в docs_path изменилась (кол-во, имена, mtime, размер)
        - Есть файлы-источники, а в манифесте было 0 (исторический stub)
        - Файлы в docs новее chroma.sqlite3
        """
        sqlite_path = os.path.join(self.vector_store_path, "chroma.sqlite3")
        if not os.path.exists(self.vector_store_path):
            return True
        if not os.path.isfile(sqlite_path):
            return True

        manifest = self._load_manifest()
        current_files = self._gather_source_files()
        current_sig = self._compute_signature(current_files)

        # Если нет файлов в docs вообще — не пересобираем (избегаем пустого индекса)
        if len(current_files) == 0:
            # Но если раньше было что-то (в манифесте >0) — не пересобираем, оставляем существующее
            return False

        if manifest is None:
            print("[manifest] отсутствует, требуется пересборка")
            return True
        old_sig = manifest.get("signature")
        old_files = manifest.get("files", [])
        if old_sig != current_sig:
            print("[manifest] обнаружены изменения в файлах документов -> пересборка")
            return True

        # Дополнительная проверка на свежесть: если любой файл новее sqlite
        try:
            index_mtime = os.path.getmtime(sqlite_path)
            for f in current_files:
                full = os.path.join(self.docs_path, f["rel"])
                if os.path.getmtime(full) > index_mtime:
                    print(f"[manifest] файл {f['rel']} новее индекса")
                    return True
        except Exception:
            pass

        return False

    # ================== КОНЕЦ БЛОКА МАНИФЕСТА ==================
